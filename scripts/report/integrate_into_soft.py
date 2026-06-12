#!/usr/bin/env python3
"""Integrate our additions INTO the team's base report (soft.docx), in its style.

Loads soft.docx verbatim (100% preserved) and appends our new chapters — Data
Architecture / Synchronization / BI, and Extended System Testing — formatted to
match the base document (chapter 20pt UPPERCASE, subsection 16pt, body 11pt,
manual table borders, mermaid figures). Output: docs/report/NextStay_Report_Integrated.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from scripts.report.build_report import render_mermaid

ROOT = Path(__file__).resolve().parents[2]
SOFT = ROOT / "soft.docx"
OUT = ROOT / "docs" / "report" / "NextStay_Report_Integrated.docx"

_tbl_no = [0]


def chapter(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.size = Pt(20)
    r.bold = False


def subsection(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.bold = False


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph()
        r = p.add_run(f"•  {it}")
        r.font.size = Pt(11)


def _borders(table):
    el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "999999")
        el.append(e)
    table._tbl.tblPr.append(el)


def table(doc, headers, rows, caption=None):
    if caption:
        _tbl_no[0] += 1
        cp = doc.add_paragraph()
        cr = cp.add_run(f"Table {_tbl_no[0]} — {caption}")
        cr.italic = True
        cr.font.size = Pt(9)
    t = doc.add_table(rows=1, cols=len(headers))
    _borders(t)
    for i, h in enumerate(headers):
        rr = t.rows[0].cells[i].paragraphs[0].add_run(h)
        rr.bold = True
        rr.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(10)
    doc.add_paragraph()


def figure(doc, code, caption):
    png = render_mermaid(code)
    if png:
        doc.add_picture(str(png), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    doc.add_paragraph()


def build():
    from scripts.report.integrate_content import add_chapters

    doc = Document(str(SOFT))
    helpers = dict(chapter=chapter, subsection=subsection, body=body, bullets=bullets, table=table, figure=figure)
    add_chapters(doc, helpers)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Integrated report written: {OUT}")


if __name__ == "__main__":
    import sys

    sys.path.insert(0, str(ROOT))
    build()
