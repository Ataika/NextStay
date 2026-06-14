#!/usr/bin/env python3
"""Generate the NextStay Software Engineering report as a .docx.

- Mermaid diagrams are rendered to PNG via mermaid.ink (no local mermaid-cli needed)
  and embedded as figures. Rendered PNGs are cached under scripts/report/figures/.
- Pass an existing .docx as a style template to inherit its look:
      python scripts/report/build_report.py --template path/to/their.docx
  Without a template, a clean default style is used.

Output: docs/report/NextStay_Report.docx
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import os
import sys
import zlib
from pathlib import Path

import httpx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
FIG_DIR = Path(__file__).resolve().parent / "figures"
OUT = Path(os.getenv("REPORT_OUT", str(ROOT / "docs" / "report" / "NextStay_Report.docx")))

ACCENT = RGBColor(0x2B, 0x6C, 0xB0)


# --------------------------------------------------------------------------- #
# Mermaid -> PNG (via mermaid.ink)
# --------------------------------------------------------------------------- #
def render_mermaid(code: str) -> Path | None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    key = hashlib.sha1(code.encode("utf-8")).hexdigest()[:16]
    png = FIG_DIR / f"{key}.png"
    if png.exists():
        return png
    # mermaid.ink "pako" format: zlib-deflate + base64url of a JSON envelope
    payload = f'{{"code":{_json_str(code)},"mermaid":{{"theme":"neutral"}}}}'
    deflated = zlib.compress(payload.encode("utf-8"), 9)
    b64 = base64.urlsafe_b64encode(deflated).decode("ascii")
    url = f"https://mermaid.ink/img/pako:{b64}?type=png&bgColor=ffffff"
    try:
        r = httpx.get(url, timeout=30.0, follow_redirects=True)
        if r.status_code == 200 and r.content[:4] == b"\x89PNG":
            png.write_bytes(r.content)
            return png
        print(f"  ! mermaid.ink {r.status_code} for diagram {key}")
    except Exception as exc:  # pragma: no cover
        print(f"  ! mermaid render failed: {exc}")
    return None


def _json_str(s: str) -> str:
    import json

    return json.dumps(s)


def _add_table_borders(table):
    """Add single black gridlines to a table whose document lacks a table style."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tbl_pr.append(borders)


# --------------------------------------------------------------------------- #
# docx helpers
# --------------------------------------------------------------------------- #
class Report:
    def __init__(self, template: str | None):
        self._table_no = 0
        self.doc = Document(template) if template else Document()
        if template:
            self._clear_body()  # keep the template's styles/theme/page setup, drop its content
        else:
            normal = self.doc.styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(11)

    def _clear_body(self):
        from docx.oxml.ns import qn

        body = self.doc.element.body
        for child in list(body):
            if child.tag == qn("w:sectPr"):
                continue  # preserve section/page configuration
            body.remove(child)

    def title_page(self):
        d = self.doc
        for _ in range(6):
            d.add_paragraph()
        t = d.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run("NextStay")
        run.bold = True
        run.font.size = Pt(40)
        run.font.color.rgb = ACCENT
        sub = d.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = sub.add_run("Hotel Property Management System")
        r2.font.size = Pt(18)
        sub2 = d.add_paragraph()
        sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub2.add_run("Software Engineering Project Report").font.size = Pt(13)
        for _ in range(10):
            d.add_paragraph()
        team = d.add_paragraph()
        team.alignment = WD_ALIGN_PARAGRAPH.CENTER
        team.add_run("Team: Dair (Backend) · Turat (Frontend) · Atai (Data Architecture & BI)").font.size = Pt(11)
        d.add_page_break()

    def toc(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        self.doc.add_heading("Table of Contents", level=1)
        para = self.doc.add_paragraph()
        run = para.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-2" \\h \\z \\u'
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "Right-click → Update Field to build the contents."
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in (begin, instr, sep, placeholder, end):
            run._r.append(el)
        self.doc.add_page_break()

    def h1(self, text: str):
        self.doc.add_heading(text, level=1)

    def h2(self, text: str):
        self.doc.add_heading(text, level=2)

    def h3(self, text: str):
        self.doc.add_heading(text, level=3)

    def p(self, text: str):
        self.doc.add_paragraph(text)

    def bullets(self, items: list[str]):
        for it in items:
            try:
                self.doc.add_paragraph(it, style="List Bullet")
            except KeyError:
                self.doc.add_paragraph(f"•  {it}")  # template lacks the List Bullet style

    def table(self, headers: list[str], rows: list[list[str]], caption: str | None = None):
        if caption:
            self._table_no += 1
            cap = self.doc.add_paragraph()
            cr = cap.add_run(f"Table {self._table_no} — {caption}")
            cr.italic = True
            cr.font.size = Pt(9)
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        styled = False
        for style_name in ("Light Grid Accent 1", "Table Grid"):
            try:
                t.style = style_name
                styled = True
                break
            except KeyError:
                continue  # style not defined in this (template) document — try next
        if not styled:
            _add_table_borders(t)  # fallback so the table still has visible gridlines
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        self.doc.add_paragraph()

    def figure(self, mermaid_code: str, caption: str):
        png = render_mermaid(mermaid_code)
        if png:
            self.doc.add_picture(str(png), width=Inches(6.2))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # fallback: embed the mermaid source as monospace so the figure is not lost
            pre = self.doc.add_paragraph()
            run = pre.add_run(mermaid_code)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        self.doc.add_paragraph()

    def image(self, path, caption: str):
        from pathlib import Path as _P

        if not _P(path).exists():
            ph = self.doc.add_paragraph()
            pr = ph.add_run(f"[screenshot pending: {caption}]")
            pr.italic = True
            return
        self.doc.add_picture(str(path), width=Inches(6.2))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        self.doc.add_paragraph()

    def code(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)

    def save(self):
        OUT.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(OUT))


# --------------------------------------------------------------------------- #
# Content
# --------------------------------------------------------------------------- #
def build(template: str | None):
    from scripts.report.content import write_all  # local import to keep this file lean

    rep = Report(template)
    rep.title_page()
    rep.toc()
    write_all(rep)
    rep.save()
    print(f"Report written: {OUT}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", default=os.getenv("REPORT_TEMPLATE"))
    args = ap.parse_args()
    sys.path.insert(0, str(ROOT))
    build(args.template)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
